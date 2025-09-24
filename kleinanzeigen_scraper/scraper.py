"""Core scraper functionality for extracting Kleinanzeigen listing links.

This module exposes helper functions used by the Streamlit app. They are
designed to be imported independently, so unit tests can run without
requiring Streamlit. All network calls use the `requests` library and a
polite User‑Agent string.
"""

from __future__ import annotations

import math
import re
from urllib.parse import urljoin

import pandas as pd
import requests
from bs4 import BeautifulSoup
# Note: python-docx might not be installed in some contexts (e.g. tests). We
# import Document lazily inside create_download_file to avoid mandatory
# dependency during parsing and scraping.



def parse_listing_links(html: str) -> list[str]:
    """Extract unique Kleinanzeigen listing URLs from a page's HTML.

    The seller overview pages contain many anchor tags, but we only want the
    individual listing links. Listing URLs include the sub‑path `/s-anzeige/` and
    end with an ID pattern (e.g. `.../2801821674-223-8242`). This helper
    identifies those anchors and returns absolute URLs.

    Args:
        html: The raw HTML content of a seller page.

    Returns:
        A list of absolute listing URLs (no duplicates) in the order they
        appear.
    """
    soup = BeautifulSoup(html, "html.parser")
    links: list[str] = []
    seen: set[str] = set()
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if not href:
            continue
        # Accept only listing links containing the Kleinanzeigen listing path.
        if "/s-anzeige/" in href:
            # Ensure full URL
            full_url = href
            if not href.startswith("http"):
                full_url = urljoin("https://www.kleinanzeigen.de", href)
            # Only record unique entries
            if full_url not in seen:
                seen.add(full_url)
                links.append(full_url)
    return links


def extract_total_count(html: str) -> int | None:
    """Attempt to extract the number of active ads from the seller page.

    The number of active ads appears near the seller description as a German
    phrase like "1489 Anzeigen online". This function uses a regular
    expression to locate such a number.

    Args:
        html: The HTML content of the seller page.

    Returns:
        The integer count of active ads or None if not found.
    """
    match = re.search(r"(\d{1,5})\s+Anzeigen\s+online", html)
    if match:
        try:
            return int(match.group(1))
        except ValueError:
            return None
    return None


def scrape_seller_listings(base_url: str) -> list[str]:
    """Scrape all listing URLs from a Kleinanzeigen seller profile.

    This function iterates through paginated seller pages using the German
    `seite` query parameter. It stops once either the expected number of
    pages has been processed or no new listings are found. A polite User‑Agent
    header helps to avoid simple blocks.

    Args:
        base_url: The seller's profile URL (e.g.
            "https://www.kleinanzeigen.de/pro/ff-wheels-by-felgenforum"). The
            function will strip any existing query parameters and build new
            ones as needed.

    Returns:
        A list of all unique listing URLs published by the seller.
    """
    # Remove any existing query parameters from the URL
    if "?" in base_url:
        base_url = base_url.split("?")[0]

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/115.0 Safari/537.36"
        )
    }

    session = requests.Session()
    session.headers.update(headers)

    # Fetch the first page to determine total count and gather links
    resp = session.get(base_url, timeout=15)
    resp.raise_for_status()
    first_html = resp.text
    all_links: list[str] = parse_listing_links(first_html)
    total_ads = extract_total_count(first_html)

    # Estimate number of pages (25 ads per page) if we know the total count.
    num_pages = None
    if total_ads:
        num_pages = math.ceil(total_ads / 25)

    # Loop through subsequent pages until no new links or page limit is reached.
    page = 2
    while True:
        if num_pages and page > num_pages:
            break
        page_url = f"{base_url}?seite={page}"
        r = session.get(page_url, timeout=15)
        # Break if we receive a non‑OK status
        if r.status_code != 200:
            break
        links = parse_listing_links(r.text)
        # Stop if no new links are found
        new_links = [url for url in links if url not in all_links]
        if not new_links:
            break
        all_links.extend(new_links)
        page += 1
    return all_links


def create_download_file(links: list[str], file_type: str) -> tuple[str, bytes]:
    """Generate a downloadable file from the list of links.

    Supported file types: 'txt', 'xlsx', 'csv', 'docx'.

    Args:
        links: List of URLs to save.
        file_type: Desired format extension.

    Returns:
        A tuple (mime_type, file_bytes) suitable for Streamlit download.
    """
    if file_type == "txt":
        content = "\n".join(links)
        return "text/plain", content.encode("utf-8")
    elif file_type == "csv":
        df = pd.DataFrame({"Listing URL": links})
        csv_str = df.to_csv(index=False)
        return "text/csv", csv_str.encode("utf-8")
    elif file_type == "xlsx":
        df = pd.DataFrame({"Listing URL": links})
        import io
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False, engine="openpyxl")
        buffer.seek(0)
        return (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            buffer.getvalue(),
        )
    elif file_type == "docx":
        # Import python-docx lazily because it might not be installed in all
        # environments (e.g. during unit tests). If missing, raise an
        # informative error.
        try:
            from docx import Document  # type: ignore
        except ImportError:
            raise ImportError(
                "python-docx ist nicht installiert. Bitte fügen Sie es zu Ihren "
                "Abhängigkeiten hinzu, wenn Sie das DOCX-Format verwenden möchten."
            )
        doc = Document()
        doc.add_heading("Kleinanzeigen Listing Links", level=1)
        for link in links:
            doc.add_paragraph(link, style="BodyText")
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
        )
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
